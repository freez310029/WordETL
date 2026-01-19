import io
import pandas as pd
from flask import Flask, render_template_string, request, send_file
from docx import Document

app = Flask(__name__)

# Basic HTML template
HTML_TEMPLATE = '''
<!DOCTYPE html>
<html>
<head><title>Word Table to Excel</title></head>
<body>
    <h2>Upload Word Document (.docx)</h2>
    <form method="post" action="/convert" enctype="multipart/form-data">
        <input type="file" name="file" accept=".docx" required>
        <button type="submit">Convert All Tables to Excel</button>
    </form>
</body>
</html>
'''

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        return "No file uploaded", 400
    
    file = request.files['file']
    doc = Document(file)
    
    if not doc.tables:
        return "No tables found in this document.", 400

    # 1. Create an in-memory buffer
    output = io.BytesIO()

    # 2. Use ExcelWriter to handle multiple sheets
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        for i, table in enumerate(doc.tables):
            data = []
            for row in table.rows:
                # Extract text from each cell in the row
                data.append([cell.text.strip() for cell in row.cells])
            
            if data:
                # Create DataFrame
                # Note: This assumes the first row of every table is a header
                df = pd.DataFrame(data[1:], columns=data[0])
                
                # Write to a unique sheet name (e.g., Table_1, Table_2)
                sheet_name = f'Table_{i+1}'
                df.to_excel(writer, sheet_name=sheet_name, index=False)

    # 3. Prepare the buffer for downloading
    output.seek(0)

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name='converted_tables.xlsx'
    )

if __name__ == '__main__':
    app.run(debug=True)