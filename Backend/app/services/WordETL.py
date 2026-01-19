import io
import pandas as pd
from docx import Document
from fastapi import HTTPException


class ConverterService:
    """Service for converting Word documents to Excel"""
    
    @staticmethod
    async def convert_word_to_excel(file_content: bytes, filename: str) -> io.BytesIO:
        """
        Convert Word document tables to Excel file
        
        Args:
            file_content: The binary content of the Word file
            filename: The name of the uploaded file
            
        Returns:
            BytesIO: Excel file as binary stream
            
        Raises:
            HTTPException: If file type is invalid or no tables found
        """
        # Validate file type
        if not filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="Only .docx files are allowed")
        
        # Read the document
        doc = Document(io.BytesIO(file_content))
        
        # Check if document has tables
        if not doc.tables:
            raise HTTPException(status_code=400, detail="No tables found in this document")
        
        # Create an in-memory buffer
        output = io.BytesIO()
        
        # Use ExcelWriter to handle multiple sheets
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            for i, table in enumerate(doc.tables):
                data = []
                for row in table.rows:
                    # Extract text from each cell in the row
                    data.append([cell.text.strip() for cell in row.cells])
                
                if data:
                    # Create DataFrame
                    # Note: This assumes the first row of every table is a header
                    df = pd.DataFrame(data[1:], columns=data[0])
                    
                    # Write to a unique sheet name (e.g., Table_1, Table_2)
                    sheet_name = f'Table_{i+1}'
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
        
        # Prepare the buffer for reading
        output.seek(0)
        
        return output